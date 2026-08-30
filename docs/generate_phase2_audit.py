from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "PharmacyOS_Phase_2_Test_and_Feature_Audit.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172529"
MUTED = "5D6A67"
PALE = "E8EEF5"
GREEN = "DDEFE8"
AMBER = "FFF0CC"
RED = "F8DEDE"
WHITE = "FFFFFF"
PAGE_WIDTH_DXA = 9360


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_numbering_definition(document: Document) -> tuple[int, int]:
    numbering = document.part.numbering_part.element
    existing_abstract = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_nums = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]

    bullet_abstract_id = max(existing_abstract, default=0) + 1
    bullet_num_id = max(existing_nums, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(bullet_abstract_id))
    bullet_multi = OxmlElement("w:multiLevelType")
    bullet_multi.set(qn("w:val"), "singleLevel")
    abstract.append(bullet_multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), "Calibri")
    run_fonts.set(qn("w:hAnsi"), "Calibri")
    run_properties.append(run_fonts)
    level.append(run_properties)
    bullet_p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    bullet_p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    bullet_p_pr.append(ind)
    level.append(bullet_p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(bullet_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(bullet_abstract_id))
    num.append(abstract_ref)

    number_abstract_id = bullet_abstract_id + 1
    number_num_id = bullet_num_id + 1
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(number_abstract_id))
    number_multi = OxmlElement("w:multiLevelType")
    number_multi.set(qn("w:val"), "singleLevel")
    abstract_num.append(number_multi)
    number_level = OxmlElement("w:lvl")
    number_level.set(qn("w:ilvl"), "0")
    for tag, value in (("start", "1"), ("numFmt", "decimal"), ("lvlText", "%1.")):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        number_level.append(node)
    number_p_pr = OxmlElement("w:pPr")
    number_tabs = OxmlElement("w:tabs")
    number_tab = OxmlElement("w:tab")
    number_tab.set(qn("w:val"), "num")
    number_tab.set(qn("w:pos"), "540")
    number_tabs.append(number_tab)
    number_p_pr.append(number_tabs)
    number_ind = OxmlElement("w:ind")
    number_ind.set(qn("w:left"), "540")
    number_ind.set(qn("w:hanging"), "270")
    number_p_pr.append(number_ind)
    number_level.append(number_p_pr)
    abstract_num.append(number_level)
    numbering.append(abstract_num)
    numbering.append(num)
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(number_num_id))
    number_ref = OxmlElement("w:abstractNumId")
    number_ref.set(qn("w:val"), str(number_abstract_id))
    number.append(number_ref)
    numbering.append(number)
    return bullet_num_id, number_num_id


def restart_numbering(document: Document, base_num_id: int) -> int:
    numbering = document.part.numbering_part.element
    base = next(node for node in numbering.findall(qn("w:num")) if int(node.get(qn("w:numId"))) == base_num_id)
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(new_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    number.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    number.append(level_override)
    numbering.append(number)
    return new_num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def add_bullet(document: Document, text: str, bullet_num_id: int, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    apply_numbering(paragraph, bullet_num_id)
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[int], status_column: int | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = value
        set_cell_shading(cell, PALE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for values in rows:
        row = table.add_row()
        keep_row_together(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if status_column == index:
                fill = GREEN if value.startswith("Built") or value.startswith("Pass") else AMBER if value.startswith("Partial") or value.startswith("Pending") else RED
                set_cell_shading(cell, fill)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    set_table_width(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_document(document: Document) -> tuple[int, int]:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color, before, after in (
        ("Title", 28, INK, 0, 16),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.style = "Table Grid"
    table.cell(0, 0).text = "PHARMACYOS  /  ENGINEERING CONTROL"
    table.cell(0, 1).text = "PHASE 2 AUDIT  ·  20 AUG 2026"
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in table.rows[0].cells:
        set_cell_shading(cell, DARK_BLUE)
        set_cell_margins(cell, top=55, start=100, bottom=55, end=100)
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    set_table_width(table, [5200, 4160])
    header.paragraphs[0]._element.getparent().remove(header.paragraphs[0]._element)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Internal implementation record  ·  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(paragraph)
    return add_numbering_definition(document)


def build() -> None:
    document = Document()
    bullet_num_id, number_num_id = configure_document(document)
    core = document.core_properties
    core.title = "PharmacyOS Phase 2 Test Plan and Feature Audit"
    core.subject = "Implementation status, gaps, and release verification plan"
    core.author = "PharmacyOS Engineering"

    document.add_heading("Phase 2 Test Plan\nand Feature Audit", 0)
    subtitle = document.add_paragraph("Operational Intelligence & Owner Assistance")
    subtitle.style = document.styles["Subtitle"]
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    subtitle.runs[0].font.size = Pt(16)
    add_table(
        document,
        ["Audit field", "Recorded value"],
        [
            ["Repository", str(ROOT)],
            ["Requirements source", "PharmacyOS Phase 2 Architecture and Implementation Prompts"],
            ["Audit date", date(2026, 8, 20).isoformat()],
            ["Verification policy", "Lightweight only: lint, strict TypeScript, production build, disposable PostgreSQL migration smoke"],
            ["Release decision", "NOT RELEASE-READY — coding foundation is substantial; listed P0 proof and integration gaps remain"],
        ],
        [2200, 7160],
    )

    document.add_heading("Executive verdict", level=1)
    p = document.add_paragraph()
    p.add_run("Outcome: ").bold = True
    p.add_run(
        "The Phase 2 backend foundation, durable jobs, core deterministic calculations, RBAC permissions, and operational UI surfaces are implemented. "
        "The repository is now materially ahead of its earlier foundation state, but it is not honest to label the complete commercial feature set finished until missing UI links, external adapters, and representative-data verification are completed."
    )
    add_bullet(document, "Verified now: zero-warning lint; strict TypeScript across all workspaces; production build; PostgreSQL 18 migration chain on a disposable database.", bullet_num_id, "Verified now:")
    add_bullet(document, "Deliberately not run now: full unit/integration/E2E/concurrency/security/performance suites, per the request to prioritize implementation first.", bullet_num_id, "Deliberately not run now:")
    add_bullet(document, "Release blockers: real FBR return adapter, complete return/procurement operator UI, backup/restore evidence, hardware QR validation, and P0 regression/concurrency proof.", bullet_num_id, "Release blockers:")

    document.add_page_break()
    document.add_heading("1. Verification actually performed", level=1)
    add_table(
        document,
        ["Gate", "Result", "Evidence / boundary"],
        [
            ["ESLint", "Pass", "npm run lint completed with zero warnings and zero errors."],
            ["Strict TypeScript", "Pass", "npm run typecheck passed for API, web, worker, config, database, and shared workspaces."],
            ["Production build", "Pass", "npm run build completed; Vite produced the production web bundle and all Node packages compiled."],
            ["Database migration smoke", "Pass", "Migrations 001–005 applied to disposable `postgres:18-alpine`; schema query returned 54 public tables."],
            ["Full automated tests", "Pending", "Not run in this implementation-first pass."],
            ["Docker Compose application smoke", "Pending", "Requires configured project secrets, seeded branch/users, and deliberate application test data."],
        ],
        [2200, 1350, 5810],
        status_column=1,
    )

    document.add_page_break()
    document.add_heading("2. Feature implementation audit", level=1)
    add_table(
        document,
        ["Feature", "Status", "Implemented now", "Remaining gap / proof"],
        [
            ["Shelf optimization", "Built / unproven", "Storage/security metadata; deterministic weekly scoring; saved reasons; one pending recommendation invariant; RBAC list/review/apply; inventory UI.", "Run representative demand scoring, storage mismatch, inactive-target, and concurrent review tests."],
            ["Expiry-loss prevention", "Partial", "Configurable policy fields; timezone-aware risk queries; acquisition-cost value at risk; idempotent work items; quarantine/action API; worker refresh; attention UI.", "UI currently displays the queue but does not expose every action/assignment control. Validate Asia/Karachi date boundaries."],
            ["Supplier price intelligence", "Partial", "Historical paid-cost derivation; discount/bonus/base-unit normalization; current quotes; audited quote entry API; comparison API.", "Dedicated quote-entry and product comparison screens are not built. Reconcile against real goods-receipt samples."],
            ["Reorder / forecasting v2", "Built / unproven", "Demand windows; stockout confidence; observed/fallback lead time; safety/coverage/MOQ/multiple math; expiry flag; idempotent daily generation; review; duplicate-safe draft PO API; queue UI.", "Draft-PO supplier/quantity selection UI is not built. Run deterministic math and retry/concurrency cases."],
            ["Budget regimen calculator", "Partial", "Exact integer/decimal calculation; complete-day maximization; pack/minimum increments; current price versions; optional audit; POS-cart UI; safety copy.", "Pharmacist verification workflow UI and checkout price-change acknowledgement are not wired end-to-end."],
            ["QR-assisted returns", "Built / unproven", "Opaque UUID token; authenticated lookup; printable token-only QR; item selection and disposition UI; request/approval/cash refund; remaining-quantity trigger; stock ledger; retry-safe operations; FBR return outbox.", "Printer/scanner hardware and concurrent partial returns are not yet validated. Non-cash refund UI and real FBR credit-note adapter remain."],
            ["Owner AI assistant", "Built / provider unproven", "Provider boundary; environment-configured model; nine whitelisted read-only tools; runtime RBAC; aggregate/sanitized facts; durable audit; database-backed rate limit; timeout/fail-closed behavior; owner UI.", "No Gemini key/model was configured or called. Mock/provider failure, authorization, and numeric-reconciliation tests remain."],
            ["Background scheduling", "Built / unproven", "Existing outbox reused; daily availability/expiry/reorder and weekly shelf jobs; dedupe keys; retry tracking; persisted run status/duration/error.", "Run worker replay, crash recovery, multi-worker lock, and resource-contention tests."],
        ],
        [1450, 1450, 3250, 3210],
        status_column=1,
    )

    document.add_heading("3. Phase 1 dependency conflicts and features not built", level=1)
    add_table(
        document,
        ["Gap", "Impact", "Required completion"],
        [
            ["Cash-session release proof", "Open, movements, reconciliation, closing, and independent variance approval are implemented in API and UI.", "Run retry, wrong-terminal, concurrent close, refund, and threshold-boundary integration cases before pilot."],
            ["Goods receipt operator UI", "Duplicate-safe partial/full goods receipt, acquisition-lot batch creation, bonus/base conversion, and stock ledger are implemented in the API.", "Build guided PO/receipt screens and reconcile representative supplier invoices through the UI."],
            ["Receipt printer/scanner validation", "An 80 mm print template renders a medium-error-correction QR containing only the opaque return token.", "Validate supported printers, paper widths, degraded prints, scanner/camera input, and token-only decoding."],
            ["Real FBR adapters", "Sandbox/disabled boundaries exist; production submission and return/credit-note behavior is not certified.", "Implement selected adapter, sandbox/certification tests, retry classification, and reconciliation."],
            ["Backup automation and restore drill", "Schema records backup runs, but no scheduled encrypted backup/restore implementation or evidence exists.", "Add deployment-specific encrypted backup job, retention, monitoring, and restore proof."],
            ["Return UI completion proof", "Lookup, selectable quantities, reason/disposition, approval, cash refund, and error states are implemented.", "Add non-cash tender controls, explicit role-handoff queue, refund receipt, and database-backed retry/concurrency proof."],
            ["Complete procurement UI", "Backend supports quotes, comparison, review, and draft creation; current UI focuses on attention/review.", "Add supplier comparison/quote entry and draft PO editing/approval screens."],
            ["Production observability", "Job runs and audit metadata exist, but metrics export/alerts and centralized support views are incomplete.", "Add metric endpoint/export, alert thresholds, structured correlation IDs, and log redaction review."],
        ],
        [2100, 3500, 3760],
    )

    document.add_page_break()
    document.add_heading("4. Required test plan", level=1)
    p = document.add_paragraph("Execution order is risk-based. Do not start broad UI polish until P0 data-integrity and authorization cases pass.")
    p.paragraph_format.keep_with_next = True

    document.add_heading("4.1 P0 — release-blocking correctness and security", level=2)
    p0_rows = [
        ["Migrations", "Reset empty PostgreSQL; apply 001–005 twice through the migration runner; verify checksums and expected failure on edited applied migration.", "Clean forward migration; second runner pass is a no-op; all constraints/indexes exist."],
        ["POS regression", "Finalize paid sale with multi-batch FEFO reservation, exact payment, retry same client request, and changed-stock race.", "One sale/invoice/token; correct batch quantities/ledger; retry returns same result."],
        ["Return concurrency", "Two sessions request the final eligible quantity at the same time; repeat scan, approval, and refund.", "Database trigger prevents over-return; one refund; original sale lines unchanged."],
        ["Shelf safety", "Cold and secured medicines against ambient/unsecured/inactive targets; target deactivated after generation.", "Unsafe target never generated/applied; stale review fails atomically."],
        ["Expiry boundary", "Expiry yesterday/today/+30/+31/+60/+61/+90 in branch timezone, including server UTC boundary.", "Buckets are exact; expired stock is unsellable; value uses acquisition cost."],
        ["Reorder math", "No history, heavy stockouts, missing lead time, MOQ, multiple rounding, near-expiry stock, repeated generation.", "Explainable quantity/confidence; one active suggestion; no automatic order."],
        ["Budget safety", "Below one day, exact budget, pack rounding, fractional units, multiple medicines, price change before checkout.", "No partial-dose result; no float drift; stale calculation is rejected/recalculated."],
        ["RBAC", "Allowed and denied identity for every new endpoint; system admin without business grants; owner AI tool re-check.", "Backend denies every unauthorized request regardless of UI visibility."],
        ["AI isolation", "Disabled provider, no key, timeout, 429, malformed response, contradictory generated number, tool failure.", "POS/health remain available; facts stay authoritative; no fabrication or secret/PII leakage."],
    ]
    add_table(document, ["Area", "Test", "Pass condition"], p0_rows, [1500, 4300, 3560])

    document.add_heading("4.2 P1 — workflow and recovery", level=2)
    p1_items = [
        "Morning inventory journey: worker refresh → attention counts → expiry review → shelf review → reorder review.",
        "Reorder-to-draft journey: compare paid cost/current quote → choose supplier/quantity → retry draft creation → approve/order without overwriting a human-edited draft.",
        "Return journey: scan printed QR → lookup → partial request → approval → disposition → refund → FBR return boundary.",
        "Worker recovery: kill during each refresh, replay outbox job, run two workers, confirm deduplication and run audit.",
        "Internet/provider loss: disconnect WAN while POS, inventory, returns, and deterministic owner reports remain usable.",
        "Database/Redis outage behavior: bounded failure, stable errors, no partial state, successful recovery after dependency returns.",
    ]
    for item in p1_items:
        add_bullet(document, item, bullet_num_id)

    document.add_heading("4.3 P2 — performance, usability, and operations", level=2)
    p2_items = [
        "Benchmark catalog/POS, expiry queue, reorder queue, supplier history, shelf scoring, and owner tools on representative pharmacy volume.",
        "Keyboard-only and small-screen review for all new screens; verify focus, disabled, loading, empty, denied, timeout, and retry states.",
        "Receipt printer/scanner matrix: supported paper widths, QR error correction, low-quality print, repeated scan, and offline LAN lookup.",
        "Backup/restore: encrypted backup containing all new durable tables; restore to clean host; compare row counts/checksums and launch app.",
        "Log/security review: secrets, tokens, prescription/health data, customer identity, payment/FBR material, and model prompts never appear in logs.",
        "Retention proof: regulated history remains queryable for configured period; cleanup jobs cannot remove sales/returns/receipts/movements." ,
    ]
    for item in p2_items:
        add_bullet(document, item, bullet_num_id)

    document.add_heading("5. Minimum representative test data", level=1)
    data_rows = [
        ["Branches/time", "At least 2 branches; Asia/Karachi boundary records around 00:00 local and UTC rollover."],
        ["Catalog", "500+ medicines; ambient/cold/secured classes; box/strip/tablet conversions; inactive and new products."],
        ["Inventory", "Multiple FEFO batches; zero/depleted; expired/today/30/31/60/61/90-day expiries; quarantined/secured stock."],
        ["Sales", "90+ days; cancelled/voided; stockout intervals; multiple counters; partial/full returns; mixed tender."],
        ["Procurement", "3+ suppliers/product; varied lead times; discounts; bonus units; quotes with validity; partial receipts."],
        ["Users", "Every seeded role plus custom least-privilege roles; explicit denied cases for owner-only/business approvals."],
        ["Failures", "Provider timeout/429/malformed; worker crash; duplicate outbox; database/Redis interruption; stale shelf/price state."],
    ]
    add_table(document, ["Dataset", "Required content"], data_rows, [1800, 7560])

    document.add_page_break()
    document.add_heading("6. Suggested verification commands", level=1)
    commands = [
        "npm run format:check",
        "npm run lint",
        "npm run typecheck",
        "npm run test",
        "npm run build",
        "powershell -ExecutionPolicy Bypass -File infra/scripts/smoke-migrations.ps1",
        "docker compose -f infra/docker/compose.yaml up --build",
    ]
    command_num_id = restart_numbering(document, number_num_id)
    for command in commands:
        paragraph = document.add_paragraph()
        apply_numbering(paragraph, command_num_id)
        run = paragraph.add_run(command)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    note = document.add_paragraph()
    note.add_run("Important: ").bold = True
    note.add_run("Record actual exit codes and outputs. Do not mark any gate passed merely because code exists or a command is planned.")

    document.add_heading("7. Release gate", level=1)
    add_table(
        document,
        ["Decision", "Condition"],
        [
            ["BLOCK", "Any authorization bypass, over-return/duplicate refund, inventory corruption, unsafe shelf placement, money/rounding error, secret/PII leakage, or failed restore."],
            ["HOLD", "Missing physical QR, production FBR path, cash-session workflow, complete return/procurement UI, or P0 automated evidence."],
            ["PILOT", "All P0 pass; critical P1 journeys pass; representative performance is acceptable; backup/restore and operations runbook are proven."],
        ],
        [1450, 7910],
    )

    document.add_heading("8. Immediate next implementation order", level=1)
    next_steps = [
        "Run database-backed cash-session and checkout retry/concurrency proof; the operational API and UI are implemented.",
        "Validate the token-only QR receipt on supported printers and scanners; browser rendering is implemented.",
        "Finish return and procurement screens over the implemented APIs.",
        "Add focused P0 unit/database/integration/concurrency tests; then run the complete repository suite.",
        "Configure the selected FBR adapter and Gemini provider only in controlled sandbox environments.",
        "Implement encrypted backup automation and produce restore evidence before any pharmacy pilot.",
    ]
    next_step_num_id = restart_numbering(document, number_num_id)
    for step in next_steps:
        paragraph = document.add_paragraph()
        apply_numbering(paragraph, next_step_num_id)
        paragraph.add_run(step)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
